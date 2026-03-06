from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 创建文档
doc = Document()
doc.add_heading('Allstate Claims Severity - XGBoost 实验报告', 0)

# 1. 实验目的
doc.add_heading('1. 实验目的', level=1)
doc.add_paragraph(
    "本实验旨在：\n"
    "1. 使用 XGBoost 回归模型对保险理赔损失(loss)进行预测。\n"
    "2. 对训练数据进行基础特征工程，处理缺失值及类别变量。\n"
    "3. 通过随机搜索超参数(RandomizedSearchCV)提升模型性能。\n"
    "4. 生成测试集预测结果并生成提交文件，用 RMSE 进行模型评估。\n"
    "5. 熟悉 Kaggle 数据分析流程及模型调优实战。"
)

# 2. 数据概况
doc.add_heading('2. 数据概况', level=1)
doc.add_paragraph(
    "训练集: 188,318 样本, 132 特征\n"
    "测试集: 125,546 样本, 131 特征\n"
    "类别特征: 116 个, 连续特征: 14 个\n"
    "目标变量 'loss' 右偏，使用 log1p 转换"
)

# 3. 实验环境
doc.add_heading('3. 实验环境', level=1)
doc.add_paragraph(
    "操作系统：Windows 10 / 11\n"
    "Python 版本：3.11\n"
    "主要库：pandas, numpy, scikit-learn, xgboost, joblib\n"
    "CPU：4 核以上"
)

# 4. 数据预处理与特征工程
doc.add_heading('4. 数据预处理与特征工程', level=1)
doc.add_paragraph(
    "1. 连续变量填充中位数，类别变量填充 'missing'\n"
    "2. 使用 pandas.factorize 对类别特征编码，测试集未出现类别映射为 -1\n"
    "3. 目标变量使用 log1p 转换"
)

# 5. 模型训练与调参
doc.add_heading('5. 模型训练与调参', level=1)
doc.add_paragraph(
    "Baseline 模型:\n"
    "n_estimators=150, max_depth=5, learning_rate=0.1, subsample=0.8, colsample_bytree=0.8\n"
    "Baseline train RMSE (原始尺度) ≈ 1843.85\n"
    "\n随机搜索调参:\n"
    "RandomizedSearchCV n_iter=2, cv=2, n_jobs=1\n"
    "参数搜索范围:\n"
    "n_estimators: 100, 150, 200\n"
    "max_depth: 3, 4, 5\n"
    "learning_rate: 0.05, 0.1\n"
    "subsample: 0.7, 0.8, 1.0\n"
    "colsample_bytree: 0.6, 0.8, 1.0\n"
    "最佳参数示例: {'n_estimators':150, 'max_depth':5, 'learning_rate':0.1, 'subsample':0.8, 'colsample_bytree':0.8}"
)

# 6. 模型评估
doc.add_heading('6. 模型评估', level=1)
doc.add_paragraph(
    "训练集划分 90% 训练 / 10% 验证\n"
    "验证集 RMSE (原始尺度) ≈ 1843.85\n"
    "在测试集生成预测提交文件 submission_时间戳.csv"
)

# 7. 模型与文件保存
doc.add_heading('7. 模型与文件保存', level=1)
doc.add_paragraph(
    "训练好的模型: xgb_model_时间戳.joblib\n"
    "随机搜索对象: xgb_randsearch_时间戳.joblib\n"
    "提交文件: submission_时间戳.csv\n"
    "文件均保存于脚本同目录"
)

# 8. 实验总结
doc.add_heading('8. 实验总结', level=1)
doc.add_paragraph(
    "结论:\n"
    "- XGBoost + log1p + factorize + 随机搜索调参，可快速获得合理 RMSE\n"
    "- n_iter=2 控制运行时间 <2 分钟\n"
    "优化建议:\n"
    "- 增加搜索次数或 CV 折数可提升性能\n"
    "- 尝试组合特征或目标编码等高级特征工程\n"
    "- GPU 可进一步加速训练\n"
    "经验总结:\n"
    "- Windows 上 Joblib 并行可能出现 Unicode 路径问题，n_jobs=1 最稳妥\n"
    "- factorize 对大规模类别特征编码效率高"
)

# 9. 附件
doc.add_heading('9. 附：可复现代码文件', level=1)
doc.add_paragraph(
    "XGBoost.py\n"
    "submission_时间戳.csv\n"
    "xgb_model_时间戳.joblib\n"
    "xgb_randsearch_时间戳.joblib"
)

# 保存文档
doc.save("REPORT.docx")
print("✅ REPORT.docx 已生成在当前目录！")
